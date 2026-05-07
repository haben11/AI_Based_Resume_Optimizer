from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from app.utils.cv_parser import parse_optimized_cv

def generate_professional_docx(markdown_content: str) -> BytesIO:
    """
    Generates a professional DOCX resume from Markdown content.
    Focuses on standard professional formatting.
    """
    data = parse_optimized_cv(markdown_content)
    doc = Document()

    # Style definitions
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run(data["full_name"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = " | ".join(data["contact_info"].values())
    contact.add_run(contact_text).font.size = Pt(9)

    # Summary
    if data["summary"]:
        doc.add_heading("PROFESSIONAL SUMMARY", level=1)
        doc.add_paragraph(data["summary"])

    # Experience
    if data["experience"]:
        doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
        for job in data["experience"]:
            p = doc.add_paragraph()
            p.add_run(job["title"]).bold = True
            for bullet in job["bullets"]:
                doc.add_paragraph(bullet, style='List Bullet')

    # Skills
    if data["skills"]:
        doc.add_heading("CORE COMPETENCIES", level=1)
        doc.add_paragraph(", ".join(data["skills"]))

    # Education
    if data["education"]:
        doc.add_heading("EDUCATION", level=1)
        for edu in data["education"]:
            doc.add_paragraph(edu)

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
